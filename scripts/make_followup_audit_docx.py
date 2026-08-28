"""Build the collaborator-facing SolvAI follow-up audit as a styled DOCX."""

from __future__ import annotations

import re
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "reviews" / "MICHAEL_BORIS_FOLLOWUP_AUDIT.md"
OUTPUT = ROOT / "reviews" / "SolvAI_Michael_Boris_Followup_Audit.docx"
REPOSITORY = "https://github.com/Scientific-Computing-Lab/SolvAI"

NAVY = "17324D"
BLUE = "2367A8"
TEAL = "16858C"
PALE_BLUE = "EAF2F8"
PALE_TEAL = "E9F5F4"
MID_GRAY = "6B7280"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(
    cell, *, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90
) -> None:
    properties = cell._tc.get_or_add_tcPr()
    margins = properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        properties.append(margins)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    properties.append(repeat)


def set_paragraph_shading(paragraph, fill: str) -> None:
    properties = paragraph._p.get_or_add_pPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    properties.append(cant_split)


def set_font(
    run, name: str, size: float, color: str | None = None, bold: bool | None = None
) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, "Aptos", 8, MID_GRAY)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_bottom_rule(paragraph, color: str = "B9C5D1") -> None:
    properties = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)
    properties.append(borders)


def build_reference_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.75)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.75)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string("202B35")
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.06

    for name in ("Body Text", "First Paragraph", "Compact"):
        if name in styles:
            style = styles[name]
            style.font.name = "Aptos"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
            style.font.size = Pt(9.5)
            style.font.color.rgb = RGBColor.from_string("202B35")
            style.paragraph_format.space_after = Pt(5)

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    title.font.size = Pt(27)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(NAVY)
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(7)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = RGBColor.from_string(TEAL)
    subtitle.paragraph_format.space_after = Pt(12)

    heading_specs = {
        "Heading 1": (17, NAVY, 16, 7),
        "Heading 2": (13.5, BLUE, 13, 5),
        "Heading 3": (10.5, TEAL, 9, 3),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = styles[name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Source Code" in styles:
        code = styles["Source Code"]
    else:
        code = styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    code.font.name = "Aptos Mono"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Mono")
    code.font.size = Pt(8.2)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Cm(0.45)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(6)

    header = section.header.paragraphs[0]
    header.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run = header.add_run("SolvAI  ·  Collaborator follow-up audit")
    set_font(run, "Aptos", 8, MID_GRAY, bold=True)
    add_bottom_rule(header)
    add_page_number(section.footer.paragraphs[0])

    # Keep the reference document minimal; Pandoc retains its styles and page setup.
    document.save(path)


def github_link(path: str, *, tree: bool = False) -> str:
    route = "tree" if tree else "blob"
    return f"[{path}]({REPOSITORY}/{route}/main/{path})"


def prepare_markdown() -> str:
    body = SOURCE.read_text()
    body = re.sub(r"^# Michael--Boris collaborator follow-up audit\n+", "", body)
    body = body.replace(
        "**Date:** 28 August 2026  \n"
        "**Scope:** evidence and submission decisions only; the frozen models, folds and\n"
        "manuscript were not changed.\n\n",
        "",
    )
    body = body.replace("### Collaborator question", "### Question")
    body = body.replace("### What the question tests", "### Why it matters")
    body = body.replace("### New analysis of frozen outputs", "### Frozen-output analysis")
    body = body.replace(
        "### Exact result and manuscript implication", "### Answer and manuscript implication"
    )
    body = body.replace("### Exact result", "### Answer")

    replacements = {
        "`scripts/analyze_collaborator_followup.py`": github_link(
            "scripts/analyze_collaborator_followup.py"
        ),
        "`results/followup_audit/`": github_link("results/followup_audit", tree=True),
        "`audits/leakage_audit.{json,csv,md}`": github_link("audits/leakage_audit.md"),
        "`audits/confirmatory/standardized_exclusion_records.csv`": github_link(
            "audits/confirmatory/standardized_exclusion_records.csv"
        ),
        "`audits/confirmatory/standardized_exclusion_refit_verification.json`": github_link(
            "audits/confirmatory/standardized_exclusion_refit_verification.json"
        ),
        "`audits/confirmatory/chemical_identity_matches.csv`": github_link(
            "audits/confirmatory/chemical_identity_matches.csv"
        ),
        "`results/followup_audit/source_block_summary.csv`": github_link(
            "results/followup_audit/source_block_summary.csv"
        ),
        "`results/followup_audit/teacher_fidelity_summary.csv`": github_link(
            "results/followup_audit/teacher_fidelity_summary.csv"
        ),
        "`results/paper_metrics.json`": github_link("results/paper_metrics.json"),
        "`results/followup_audit/hard_case_audit.csv`": github_link(
            "results/followup_audit/hard_case_audit.csv"
        ),
        "`hard_case_group_summary.csv`": github_link(
            "results/followup_audit/hard_case_group_summary.csv"
        ),
        "`reports/LAMBDA_RESPONSE_EXPERIMENT.md`": github_link(
            "reports/LAMBDA_RESPONSE_EXPERIMENT.md"
        ),
        "`reviews/NOVELTY_POSITIONING_AUDIT.md`": github_link(
            "reviews/NOVELTY_POSITIONING_AUDIT.md"
        ),
        "`reviews/NOVELTY_EDITORIAL_CHECK.md`": github_link("reviews/NOVELTY_EDITORIAL_CHECK.md"),
    }
    for source, replacement in replacements.items():
        body = body.replace(source, replacement)

    evidence_note = (
        "**Evidence trail.** Blue file links open the exact frozen artifact in the "
        f"[canonical SolvAI repository]({REPOSITORY}). Quantitative follow-up tables "
        "were derived without model fitting or model selection. The frozen models, "
        "folds and manuscript were not changed.\n\n"
    )
    return (
        "---\n"
        'title: "SolvAI collaborator follow-up audit"\n'
        'subtitle: "Answers to the scientific questions raised by Michael Levitt and Boris"\n'
        'date: "28 August 2026"\n'
        "---\n\n"
        f"{evidence_note}\n{body}"
    )


def style_output(path: Path) -> None:
    document = Document(path)
    document.core_properties.title = "SolvAI collaborator follow-up audit"
    document.core_properties.subject = (
        "Evidence review and Nature Communications submission decisions"
    )
    document.core_properties.keywords = "SolvAI, hydration free energy, response distillation"
    document.core_properties.comments = "Generated from frozen repository evidence."

    # Add a restrained visual divider below the subtitle/date block.
    for paragraph in document.paragraphs[:6]:
        if paragraph.style.name == "Title":
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name in {"Subtitle", "Date"}:
            paragraph.paragraph_format.keep_with_next = True

    for paragraph in document.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            paragraph.paragraph_format.keep_with_next = True
        if paragraph.style.name == "Block Text":
            paragraph.paragraph_format.left_indent = Cm(0.55)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(7)
            properties = paragraph._p.get_or_add_pPr()
            borders = OxmlElement("w:pBdr")
            left = OxmlElement("w:left")
            left.set(qn("w:val"), "single")
            left.set(qn("w:sz"), "18")
            left.set(qn("w:space"), "8")
            left.set(qn("w:color"), TEAL)
            borders.append(left)
            properties.append(borders)
        if paragraph.text.startswith("Evidence trail."):
            paragraph.paragraph_format.left_indent = Cm(0.35)
            paragraph.paragraph_format.right_indent = Cm(0.35)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(10)
            set_paragraph_shading(paragraph, PALE_TEAL)
        if paragraph.text.startswith("Recommendation:"):
            paragraph.paragraph_format.left_indent = Cm(0.3)
            paragraph.paragraph_format.right_indent = Cm(0.2)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(8)
            set_paragraph_shading(paragraph, LIGHT_GRAY)

    for paragraph in document.paragraphs[:5]:
        if paragraph.text == "28 August 2026":
            add_bottom_rule(paragraph, TEAL)

    table_widths_cm = {
        0: [3.0, 1.1, 8.6, 4.3],
        1: [5.2, 7.1, 4.7],
        2: [5.3, 1.3, 3.6, 3.4, 3.4],
        3: [5.0, 12.0],
        4: [3.2, 4.6, 9.2],
        5: [0.8, 2.3, 4.2, 2.1, 2.0, 2.8, 2.0],
        6: [5.1, 3.8, 3.9, 4.2],
    }
    for table_index, table in enumerate(document.tables):
        table.autofit = False
        widths = table_widths_cm.get(table_index)
        if widths:
            grid = table._tbl.tblGrid
            while len(grid):
                grid.remove(grid[0])
            for width_cm in widths:
                column = OxmlElement("w:gridCol")
                column.set(qn("w:w"), str(Cm(width_cm).twips))
                grid.append(column)
            properties = table._tbl.tblPr
            layout = properties.find(qn("w:tblLayout"))
            if layout is None:
                layout = OxmlElement("w:tblLayout")
                properties.append(layout)
            layout.set(qn("w:type"), "fixed")
        try:
            table.style = "Light Shading Accent 1"
        except KeyError:
            table.style = "Table Grid"
        for row_index, row in enumerate(table.rows):
            prevent_row_split(row)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            if row_index == 0:
                set_repeat_header(row)
            for cell_index, cell in enumerate(row.cells):
                if widths and cell_index < len(widths):
                    cell.width = Cm(widths[cell_index])
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                if row_index == 0:
                    set_cell_shading(cell, NAVY)
                elif row_index % 2 == 0:
                    set_cell_shading(cell, PALE_BLUE if table_index != 0 else PALE_TEAL)
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_after = Pt(1.5)
                    paragraph.paragraph_format.space_before = Pt(1.5)
                    for run in paragraph.runs:
                        set_font(
                            run,
                            "Aptos",
                            7.1
                            if len(widths or []) >= 7
                            else (7.6 if len(widths or []) >= 5 else 8.2),
                            WHITE if row_index == 0 else "202B35",
                            bold=True if row_index == 0 else None,
                        )

    # Add a short repository footer link on the last page without changing content.
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12)
    add_bottom_rule(paragraph, LIGHT_GRAY)
    run = paragraph.add_run(f"Evidence repository: {REPOSITORY}")
    set_font(run, "Aptos", 8, MID_GRAY)

    document.save(path)


def main() -> None:
    with tempfile.TemporaryDirectory(prefix="solvai_docx_") as directory:
        temporary = Path(directory)
        reference = temporary / "reference.docx"
        markdown = temporary / "audit.md"
        build_reference_docx(reference)
        markdown.write_text(prepare_markdown())
        subprocess.run(
            [
                "pandoc",
                str(markdown),
                "--from=markdown+pipe_tables+tex_math_single_backslash",
                "--to=docx",
                f"--reference-doc={reference}",
                "--output",
                str(OUTPUT),
            ],
            check=True,
        )
    style_output(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
